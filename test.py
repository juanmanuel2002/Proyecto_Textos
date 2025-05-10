import os
import fitz  # PyMuPDF
import nltk
from nltk.corpus import stopwords
from nltk.probability import FreqDist
import string
import re
from docx import Document

# Descargar recursos necesarios de NLTK
nltk.download('stopwords')
nltk.download('punkt')

# Leer el contenido de un PDF
def leer_pdf(archivo):
    doc = fitz.open(archivo)
    texto = ""
    for pagina in doc:
        texto += pagina.get_text()
    return texto

# Stopwords personalizadas
def get_stop_words():
    custom_words = {
        "1", "2", "3", "4", "5", "6", "7", "8", "9", "0", "10",
        *list("abcdefghijklmnopqrstuvwxyz"), "et", "etc", "al", "al.", "et.",
        "etc.", "i.e", "e.g", "e.g.", "i.e.", "nm", "0O", "o0", "doi", "lm",
        *[f"e{i}" for i in range(10)], "et al", "et al.", "et. al", "et. al."
    }
    return set(stopwords.words('english')).union(custom_words)

STOP_WORDS = get_stop_words()

# Preprocesamiento del texto
def preprocesar_texto(texto):
    texto = texto.lower()
    texto = texto.translate(str.maketrans('', '', string.punctuation))
    palabras = nltk.word_tokenize(texto)
    palabras_filtradas = [palabra for palabra in palabras if palabra not in STOP_WORDS]
    return palabras_filtradas

# Palabras más frecuentes
def obtener_palabras_frecuentes(texto):
    palabras_filtradas = preprocesar_texto(texto)
    frecuencia = FreqDist(palabras_filtradas)
    return frecuencia

# Extraer oraciones con patrones relacionados a causas y consecuencias
def extraer_causas_consecuencias(texto):
    oraciones = nltk.sent_tokenize(texto)
    causas = []
    consecuencias = []

    patrones_causa = ["caused by", "due to", "because of", "triggered by", "induced by"]
    patrones_consecuencia = ["results in", "leads to", "causes", "contributes to", "associated with"]

    for oracion in oraciones:
        oracion_lower = oracion.lower()
        if any(pat in oracion_lower for pat in patrones_causa):
            causas.append(oracion.strip())
        if any(pat in oracion_lower for pat in patrones_consecuencia):
            consecuencias.append(oracion.strip())
    return causas, consecuencias

# Generar conclusión simple basada en palabras clave
def generar_conclusion(texto):
    texto_lower = texto.lower()
    if 'melanoma' in texto_lower and any(w in texto_lower for w in ['risk', 'dangerous', 'malignant', 'fatal', 'aggressive']):
        return "Los documentos revisados sugieren que los melanomas son peligrosos y deben tratarse con cuidado médico."
    elif 'melanoma' in texto_lower and 'benign' in texto_lower:
        return "Algunos documentos indican que los melanomas pueden ser benignos, pero generalmente se consideran una condición seria."
    else:
        return "La revisión de los textos no permite establecer con claridad si los melanomas son peligrosos o benignos."

# Limpieza de texto para exportación segura en Word
def limpiar_texto(texto):
    if texto:
        return re.sub(r'[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]', '', texto)
    return texto

# Exportar todo a Word
def exportar_todo_a_word(causas, consecuencias, conclusion, output='analisis_melanomas.docx'):
    doc = Document()
    doc.add_heading('Análisis de Documentos sobre Melanoma', 0)

    doc.add_heading('Causas Identificadas', level=1)
    for c in causas:
        doc.add_paragraph(limpiar_texto(c), style='List Bullet')

    doc.add_heading('Consecuencias Identificadas', level=1)
    for c in consecuencias:
        doc.add_paragraph(limpiar_texto(c), style='List Bullet')

    doc.add_heading('Conclusión General', level=1)
    doc.add_paragraph(limpiar_texto(conclusion))

    doc.save(output)
    print(f"\n✅ Archivo Word exportado: {output}")

# --------- PROCESO PRINCIPAL ---------

carpeta_pdfs = './Archivos'
todos_los_textos = ""

for archivo in os.listdir(carpeta_pdfs):
    if archivo.endswith('.pdf'):
        ruta = os.path.join(carpeta_pdfs, archivo)
        texto_pdf = leer_pdf(ruta)
        todos_los_textos += texto_pdf + "\n"

# Palabras frecuentes
frecuencia_palabras = obtener_palabras_frecuentes(todos_los_textos)
print("\nPalabras más frecuentes:")
print(frecuencia_palabras.most_common(20))

# Extraer causas y consecuencias
causas, consecuencias = extraer_causas_consecuencias(todos_los_textos)

# Generar conclusión automática
conclusion = generar_conclusion(todos_los_textos)

# Exportar resultados a Word
exportar_todo_a_word(causas, consecuencias, conclusion)
